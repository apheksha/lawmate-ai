import os
import io
import re
import json
import time
import uuid
import pathlib
import tempfile
import hashlib
import logging
import platform
import subprocess
import difflib
from datetime import datetime, timedelta
from typing import Dict, Any, List, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
import streamlit as st
import pandas as pd
import plotly.express as px
import bleach
import secrets
from law_embeddings import LawDatabase
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
import torch
import html
import random
from fpdf import FPDF
from docx import Document

try:
    from weasyprint import HTML
    _HAS_WEASY = True
except Exception:
    _HAS_WEASY = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    _HAS_REPORTLAB = True
except Exception:
    _HAS_REPORTLAB = False

# SQLAlchemy
from sqlalchemy import (
    create_engine, Column, Integer, String, Text, DateTime, ForeignKey, Index,Float, Boolean, delete as sa_delete
)
from sqlalchemy.orm import declarative_base
from sqlalchemy.orm import sessionmaker, relationship

# Password hashing
try:
    from argon2 import PasswordHasher
    from argon2.exceptions import VerifyMismatchError
    _HAS_ARGON2 = True
    _argon_hasher = PasswordHasher()
except Exception:
    _HAS_ARGON2 = False
    _argon_hasher = None

import bcrypt  # fallback

# OCR
try:
    import pytesseract
    _HAS_PYTESSERACT = True
except Exception:
    pytesseract = None
    _HAS_PYTESSERACT = False

# PDF -> images
try:
    from pdf2image import convert_from_bytes
    _HAS_PDF2IMAGE = True
except Exception:
    _HAS_PDF2IMAGE = False
os.environ["TOKENIZERS_PARALLELISM"] = "false"
# Transformers 
try:
    from transformers import pipeline
    import torch
    _HAS_TRANSFORMERS = True
except Exception:
    pipeline = None
    torch = None
    _HAS_TRANSFORMERS = False


# langdetect optional (auto language detection)
try:
    from langdetect import detect as langdetect_detect
    _HAS_LANGDETECT = True
except Exception:
    langdetect_detect = None
    _HAS_LANGDETECT = False

# googletrans optional fallback for translations
try:
    from googletrans import Translator as GoogleTranslator
    _HAS_GOOGLETRANS = True
except Exception:
    GoogleTranslator = None
    _HAS_GOOGLETRANS = False

# python-docx optional (DOCX export)
try:
    from docx import Document as DocxDocument
    _HAS_DOCX = True
except Exception:
    DocxDocument = None
    _HAS_DOCX = False
st.set_page_config(page_title="LawMate AI", layout="wide")
# ---------- CONFIG ----------
DB_URL = os.getenv("DATABASE_URL", "sqlite:///lawmate.db")
CACHE_PATH = pathlib.Path(os.getenv("CACHE_FILE", "lawmate_cache.json"))
HF_MODEL = os.getenv("HF_MODEL", "mrm8488/T5-base-finetuned-cuad")
TRANSLATION_BACKEND = os.getenv("TRANSLATION_BACKEND", "none")  # 'hf' | 'google' | 'none'
TRANSLATION_MODEL = os.getenv("TRANSLATION_MODEL", "Helsinki-NLP/opus-mt-en-ROMANCE")

MAX_UPLOAD_MB = int(os.getenv("MAX_UPLOAD_MB", "20"))  # MB
SESSION_TIMEOUT_MIN = int(os.getenv("SESSION_TIMEOUT_MIN", "45"))
LOGIN_RATE_LIMIT = int(os.getenv("LOGIN_RATE_LIMIT", "5"))
LOGIN_LOCK_MIN = int(os.getenv("LOGIN_LOCK_MIN", "10"))

# Logging (structured JSON logs)
LOG_FILE = os.getenv("LOG_FILE", "lawmate_events.log")
logger = logging.getLogger("lawmate")
logger.setLevel(logging.INFO)
if not logger.handlers:
    fh = logging.FileHandler(LOG_FILE)
    fh.setFormatter(logging.Formatter('%(message)s'))
    logger.addHandler(fh)

# ---------- DB SETUP ----------
Base = declarative_base()
engine = create_engine(DB_URL, connect_args={"check_same_thread": False} if DB_URL.startswith("sqlite") else {})
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)

# Models
class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True)
    username = Column(String, unique=True, index=True, nullable=False)
    password_hash = Column(String, nullable=False)
    sec_question = Column(String, nullable=True)
    sec_answer_hash = Column(String, nullable=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    failed_logins = Column(Integer, default=0)
    locked_until = Column(DateTime, nullable=True)
    contracts = relationship("Contract", back_populates="owner", cascade="all, delete-orphan")

class Contract(Base):
    __tablename__ = "contracts"
    id = Column(Integer, primary_key=True)
    file_name = Column(String, nullable=False, index=True)
    text = Column(Text, nullable=False)
    summary = Column(Text, default="")
    created_at = Column(DateTime, default=datetime.utcnow)
    owner_id = Column(Integer, ForeignKey("users.id"), nullable=True)
    parent_id = Column(Integer, ForeignKey("contracts.id"), nullable=True)  # versioning
    owner = relationship("User", back_populates="contracts")
    qas = relationship("QACache", back_populates="contract", cascade="all, delete-orphan")

class QACache(Base):
    __tablename__ = "qa_cache"
    id = Column(Integer, primary_key=True)
    contract_id = Column(Integer, ForeignKey("contracts.id"))
    question = Column(Text)
    answer = Column(Text)
    risk_info = Column(Text)  # JSON
    created_at = Column(DateTime, default=datetime.utcnow)
    contract = relationship("Contract", back_populates="qas")

class OCRCache(Base):
    __tablename__ = "ocr_cache"
    id = Column(Integer, primary_key=True)
    file_hash = Column(String, unique=True, index=True, nullable=False)
    lang = Column(String, nullable=True)
    text = Column(Text, nullable=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    # optional: link to contract if desired later
    contract_id = Column(Integer, ForeignKey("contracts.id"), nullable=True)

Index("ix_qacache_question", QACache.question)
Index("ix_contract_file_name", Contract.file_name)
Index("ix_ocrcache_hash", OCRCache.file_hash)

Base.metadata.create_all(bind=engine, checkfirst=True)

# ---------- Compatibility shim ----------
if not hasattr(st, "rerun"):
    if hasattr(st, "experimental_rerun"):
        st.rerun = st.experimental_rerun
    else:
        def _no_rerun():
            raise RuntimeError("Streamlit rerun not available.")
        st.rerun = _no_rerun

# ---------- Helpers ----------
def now_utc() -> datetime:
    return datetime.utcnow()

def log_event(event: str, **data):
    """Write structured JSON event to the log file"""
    payload = {"ts": now_utc().isoformat(), "event": event, **data}
    try:
        logger.info(json.dumps(payload))
    except Exception:
        pass

def sanitize_filename(name: str) -> str:
    name = os.path.basename(name)
    return re.sub(r"[^\w\-_\. ]", "_", name)[:200]

def file_hash_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()

def generate_hash(text: str, question: str) -> str:
    h = hashlib.sha256()
    h.update(text.encode("utf-8", errors="ignore"))
    h.update(b"\x00")
    h.update(question.encode("utf-8", errors="ignore"))
    return h.hexdigest()

# ---------- Session helpers ----------
def refresh_session_time():
    st.session_state["last_activity"] = now_utc().isoformat()

def is_session_expired() -> bool:
    last = st.session_state.get("last_activity")
    if not last:
        return False
    try:
        last_dt = datetime.fromisoformat(last) if isinstance(last, str) else last
    except Exception:
        return False
    return (now_utc() - last_dt).total_seconds() > SESSION_TIMEOUT_MIN * 60

# ---------- Cache (JSON) ----------
def load_cache(path: pathlib.Path = CACHE_PATH) -> Dict[str, Any]:
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception as e:
            try:
                logger.info(json.dumps({"ts": now_utc().isoformat(), "event": "cache_load_failed", "error": str(e)}))
            except Exception:
                logger.info(f"cache_load_failed: {e}")
            return {}
    return {}

def save_cache(cache: Dict[str, Any], path: pathlib.Path = CACHE_PATH):
    tmp = path.with_suffix(".tmp")
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)

if 'cache' not in st.session_state:
    st.session_state.cache = load_cache()

# ---------- Password hashing----------
def hash_password(password: str) -> str:
    if _HAS_ARGON2 and _argon_hasher:
        try:
            return _argon_hasher.hash(password)
        except Exception:
            pass
    # fallback bcrypt
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode()
def password_valid_policy(password: str) -> tuple[bool, str]:
    if len(password) < 8:
        return False, "Password must be at least 8 characters."
    if not re.search(r"[A-Z]", password):
        return False, "Password must contain an uppercase letter."
    if not re.search(r"[a-z]", password):
        return False, "Password must contain a lowercase letter."
    if not re.search(r"\d", password):
        return False, "Password must contain a digit."
    return True, "OK"

def regenerate_session_token():
    token = secrets.token_hex(32)
    st.session_state.session_token = token
    return token

def verify_password(password: str, hashed: str) -> bool:
    if _HAS_ARGON2 and _argon_hasher:
        try:
            _argon_hasher.verify(hashed, password)
            return True
        except Exception:
            try:
                return bcrypt.checkpw(password.encode("utf-8"), hashed.encode())
            except Exception:
                return False
    else:
        try:
            return bcrypt.checkpw(password.encode("utf-8"), hashed.encode())
        except Exception:
            return False

# password reset helpers (security question)
def hash_answer(answer: str) -> str:
    return hash_password(answer.strip().lower())

def verify_answer(ans_plain: str, ans_hash: str) -> bool:
    return verify_password(ans_plain.strip().lower(), ans_hash)

# ---------- OCR helpers (with DB-persistent cache + parallel) ----------
LANG_CODE_MAP = {"English": "eng", "Hindi": "hin", "French": "fra"}

def tesseract_languages_available() -> List[str]:
    if not _HAS_PYTESSERACT:
        return []
    try:
        return pytesseract.get_languages(config='')  # type: ignore
    except Exception:
        try:
            return pytesseract.get_languages()  # type: ignore
        except Exception:
            return []

def ocr_image(img, lang="eng") -> str:
    if not _HAS_PYTESSERACT:
        return ""
    try:
        return pytesseract.image_to_string(img, lang=lang)
    except Exception:
        try:
            return pytesseract.image_to_string(img)
        except Exception:
            return ""

def ocr_images_parallel(images, lang="eng", max_workers: int = 4):
    text_pages = []
    if not images:
        return ""
    with ThreadPoolExecutor(max_workers=min(max_workers, len(images))) as ex:
        futures = [ex.submit(ocr_image, img, lang) for img in images]
        for f in as_completed(futures):
            try:
                text_pages.append(f.result())
            except Exception:
                text_pages.append("")
    return "\n".join(text_pages)

def extract_pdf_text(file_stream: io.BytesIO, lang: str = "eng") -> str:
    # Try PyPDF2 text extraction 
    try:
        from PyPDF2 import PdfReader
        file_stream.seek(0)
        reader = PdfReader(file_stream)
        text = ""
        for p in reader.pages:
            try:
                page_text = p.extract_text()
            except Exception:
                page_text = None
            if page_text:
                text += page_text + "\n"
        if text.strip():
            return text
    except Exception:
        pass

    if not _HAS_PDF2IMAGE:
        log_event("pdf2image_missing", message="pdf2image not installed or Poppler missing")
        return ""

    try:
        file_stream.seek(0)
        images = convert_from_bytes(file_stream.read())
    except Exception as e:
        log_event("pdf2image_failed", error=str(e))
        return ""

    try:
        return ocr_images_parallel(images, lang=lang)
    except Exception as e:
        log_event("ocr_parallel_failed", error=str(e))
        text = ""
        for img in images:
            text += ocr_image(img, lang=lang) + "\n"
        return text

def extract_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return file_bytes.decode("latin-1", errors="ignore")

# DB-backed OCR cache: checks DB first, then session cache, then runs OCR and saves to DB
def ocr_cached(file_bytes: bytes, lang_code: str = "eng", persist_to_db: bool = True) -> str:
    file_h = file_hash_bytes(file_bytes + lang_code.encode())
    # check in-memory session cache
    if 'ocr_cache' not in st.session_state:
        st.session_state.ocr_cache = {}
    if file_h in st.session_state.ocr_cache:
        return st.session_state.ocr_cache[file_h]

    # check DB cache
    try:
        with SessionLocal() as db:
            db_entry = db.query(OCRCache).filter_by(file_hash=file_h).first()
            if db_entry and db_entry.text:
                st.session_state.ocr_cache[file_h] = db_entry.text
                return db_entry.text
    except Exception as e:
        log_event("ocr_db_read_failed", error=str(e))

    # run OCR/extraction
    text = ""
    try:
        # heuristics: if bytes start with %PDF, treat as PDF
        if file_bytes[:4] == b"%PDF":
            text = extract_pdf_text(io.BytesIO(file_bytes), lang=lang_code)
        else:
            text = extract_txt(file_bytes)
    except Exception as e:
        log_event("ocr_failed", error=str(e))
        text = ""

    # update caches
    st.session_state.ocr_cache[file_h] = text
    try:
        st.session_state.cache.setdefault("ocr", {})[file_h] = {"text_preview": text[:200], "ts": now_utc().isoformat()}
        save_cache(st.session_state.cache)
    except Exception:
        pass

    if persist_to_db:
        try:
            with SessionLocal() as db:
                db_entry = db.query(OCRCache).filter_by(file_hash=file_h).first()
                if not db_entry:
                    db_entry = OCRCache(file_hash=file_h, lang=lang_code, text=text, created_at=now_utc())
                    db.add(db_entry)
                else:
                    db_entry.text = text
                    db_entry.lang = lang_code
                    db_entry.created_at = now_utc()
                db.commit()
        except Exception as e:
            log_event("ocr_db_write_failed", error=str(e))

    return text

# ---------- Risk scoring, highlighting & analytics ----------
RISK_PATTERNS = [
    (r"penalt(y|ies|amount|fee)", "High", 3),
    (r"terminate(ion|d)? without (notice|cause|consent)", "High", 3),
    (r"indemnif(y|ication)", "High", 2),
    (r"shall\s+not|must\s+not", "High", 2),
    (r"shall\b", "Medium", 1),
    (r"may\s+not|subject to", "Medium", 1),
    (r"confidential", "Medium", 1),
    (r"warrant(y|ies)?", "Low", 1),
]
NEGATION_WORDS = {"not", "no", "without", "never"}

def compute_risk(clause: str) -> Dict[str, Any]:
    score = 0
    reasons = []
    ctext = clause.lower()
    for pattern, label, weight in RISK_PATTERNS:
        for m in re.finditer(pattern, ctext):
            start = max(0, m.start()-60)
            end = min(len(ctext), m.end()+20)
            window = ctext[start:end]
            neg = any(re.search(rf"\b{re.escape(n)}\b", window) for n in NEGATION_WORDS)
            if neg:
                reasons.append(f"Found '{m.group(0)}' but negated in context.")
                score -= weight
            else:
                reasons.append(f"Matched '{m.group(0)}' -> {label}")
                score += weight
    if score >= 3:
        level = "High"
    elif score >= 1:
        level = "Medium"
    else:
        level = "Low"
    return {"score": score, "level": level, "reasons": reasons}

def highlight_text_with_risks(text: str) -> str:
    """Return sanitized HTML with risky keywords color-coded inline."""
    keywords = []
    for pat, label, weight in RISK_PATTERNS:
        w = re.sub(r"\\b|\\s|\(|\)|\||\[.*?\]|\?.*", "", pat)
        w = w.replace("(?", "").strip()
        if w:
            keywords.append((w.lower(), label))
    keywords = sorted(list({(k,l) for k,l in keywords}), key=lambda x: -len(x[0]))
    safe = bleach.clean(text, strip=True)
    html = bleach.linkify(bleach.clean(safe))
    for kw, label in keywords:
        if not kw:
            continue
        color = {"High": "#ff4b4b", "Medium": "#ffb04b", "Low": "#7bd389"}.get(label, "#ffd")
        pattern = re.compile(re.escape(kw), flags=re.IGNORECASE)
        html = pattern.sub(lambda m: f"<span style='background:{color};padding:0 2px;border-radius:3px'>{m.group(0)}</span>", html)
    return html
def highlight_search_term(text: str, search_term: str) -> str:
    """Highlight search terms safely before risk highlighting."""
    if not search_term:
        return text
    # Regex only on plain text
    return re.sub(
        re.escape(search_term),
        lambda m: f"<<HIGHLIGHT>>{m.group(0)}<</HIGHLIGHT>>",
        text,
        flags=re.IGNORECASE
    )

def render_preview_with_highlights(text: str, search_term: str) -> str:
    # Step 1: highlight search term on plain text
    marked_text = highlight_search_term(text, search_term)

    # Step 2: run your risk highlighter
    html = highlight_text_with_risks(marked_text)

    # Step 3: replace placeholders with safe <mark> tags
    html = html.replace("<<HIGHLIGHT>>", "<mark>").replace("<</HIGHLIGHT>>", "</mark>")
    return html

# ---------- Model loading & fallback ----------
@st.cache_resource(show_spinner=False)
def load_qa_pipeline(model_name=HF_MODEL):
    """Load QA or text2text pipeline safely, with fallback."""
    if not _HAS_TRANSFORMERS or pipeline is None:
        return None
    try:
        device = 0 if torch and torch.cuda.is_available() else -1
        return pipeline("question-answering", model=model_name, device=device)
    except Exception:
        try:
            return pipeline("text2text-generation", model=model_name, device=device)
        except Exception as e:
            log_event("hf_model_load_failed", error=str(e))
            return None

qa_pipeline = load_qa_pipeline()

def chunk_text(text: str, max_chars: int = 3000) -> List[str]:
    parts = []
    paras = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
    if not paras:
        paras = re.split(r'(?<=[.?!]\s)', text)
    cur = ""
    for p in paras:
        if len(cur) + len(p) <= max_chars:
            cur += (p + "\n\n")
        else:
            if cur:
                parts.append(cur.strip())
            if len(p) > max_chars:
                for i in range(0, len(p), max_chars):
                    parts.append(p[i:i+max_chars])
                cur = ""
            else:
                cur = p + "\n\n"
    if cur:
        parts.append(cur.strip())
    return parts

def query_model_with_chunking(text: str, question: str) -> str:
    MAX_CONTEXT_CHARS = 50000
    ctx = text[:MAX_CONTEXT_CHARS]
    key = generate_hash(ctx, question)
    if key in st.session_state.cache:
        log_event("cache_hit", key=key)
        return st.session_state.cache[key]['answer']
    if qa_pipeline:
        try:
            task_name = getattr(qa_pipeline, "task", "") or str(qa_pipeline)
            if "question-answering" in task_name:
                best = {"score": -1, "answer": ""}
                for chunk in chunk_text(ctx, max_chars=1500):
                    try:
                        out = qa_pipeline(question=question, context=chunk, clean_up_tokenization_spaces=True)
                        sc = float(out.get("score", 0))
                        if sc > best["score"]:
                            best = {"score": sc, "answer": out.get("answer", "")}
                    except Exception:
                        continue
                ans = best["answer"] or "No confident answer found."
                st.session_state.cache[key] = {"answer": ans, "created_at": now_utc().isoformat()}
                save_cache(st.session_state.cache)
                return ans
            else:
                chunks = chunk_text(ctx, max_chars=2500)[:3]
                prompt = "Context:\n" + "\n\n".join(chunks) + f"\n\nQuestion: {question}"
                out = qa_pipeline(prompt, max_length=512)
                if isinstance(out, list) and out:
                    ans = out[0].get('generated_text') or out[0].get('answer') or str(out[0])
                else:
                    ans = str(out)
                st.session_state.cache[key] = {"answer": ans, "created_at": now_utc().isoformat()}
                save_cache(st.session_state.cache)
                return ans
        except Exception as e:
            log_event("model_inference_failed", error=str(e))
    sentences = re.split(r'(?<=[.?!]\s)', ctx)
    question_words = [w.lower() for w in re.findall(r"\w+", question)]
    best_sentences = []
    for s in sentences:
        s_l = s.lower()
        if any(qw in s_l for qw in question_words):
            best_sentences.append(s.strip())
    ans = " ".join(best_sentences[:3]) if best_sentences else "No trustworthy automated answer available; please review the document."
    st.session_state.cache[key] = {"answer": ans, "created_at": now_utc().isoformat()}
    save_cache(st.session_state.cache)
    return ans

# ---------- Translation ----------
@st.cache_resource(show_spinner=False)
def load_translation_pipeline(model_name=TRANSLATION_MODEL):
    if not _HAS_TRANSFORMERS:
        return None
    try:
        return pipeline("translation", model=model_name)
    except Exception:
        return None

translation_pipeline = load_translation_pipeline() if TRANSLATION_BACKEND == "hf" else None

def translate_text(text: str, dest: str = "en") -> str:
    if dest == "en" or TRANSLATION_BACKEND == "none":
        return text
    if TRANSLATION_BACKEND == "hf" and translation_pipeline:
        try:
            out = translation_pipeline(text)
            if isinstance(out, list) and out:
                return out[0].get("translation_text", text)
            return str(out)
        except Exception:
            pass
    if _HAS_GOOGLETRANS:
        try:
            gt = GoogleTranslator()
            return gt.translate(text, dest=dest).text
        except Exception:
            pass
    return text

# ---------- Exports: PDF + DOCX + JSON ----------
ALLOWED_TAGS = ['div', 'span', 'br', 'strong', 'em', 'p', 'h1', 'h2', 'h3', 'ul', 'li', 'b']
ALLOWED_ATTRS = {'*': ['style']}

def sanitize_html(text: str) -> str:
    text = text.replace("\n", "<br/>")
    return bleach.clean(text, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRS, strip=True)

def export_pdf_with_metadata(answer_text: str, risk_info: List[Dict[str,Any]], username: str = "anonymous") -> Optional[str]:
    html_content = f"<h2>LawMate AI — Analysis</h2><p><strong>User:</strong> {username} &nbsp;&nbsp; <strong>Date:</strong> {now_utc().isoformat()}</p>"
    html_content += "<h3>Answer</h3><div>" + sanitize_html(answer_text) + "</div>"
    html_content += "<h3>Risk Summary</h3><ul>"
    for r in risk_info:
        reasons_str = ""
        if isinstance(r.get("reasons"), list):
            reasons_str = " — " + ", ".join(r.get("reasons"))
        elif r.get("reasons"):
            reasons_str = " — " + str(r.get("reasons"))
        html_content += f"<li><strong>{r['level']}</strong> — {sanitize_html(r['clause'])}{sanitize_html(reasons_str)}</li>"
    html_content += "</ul>"
    html_path = None
    pdf_path = None
    if _HAS_WEASY:
        try:
            with tempfile.NamedTemporaryFile(suffix=".html", delete=False, mode="w", encoding="utf-8") as fh:
                fh.write(html_content)
                html_path = fh.name
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_tmp:
                pdf_path = pdf_tmp.name
                HTML(filename=html_path).write_pdf(pdf_path)
            return pdf_path
        except Exception as e:
            log_event("weasy_failed", error=str(e))
        finally:
            if html_path and os.path.exists(html_path):
                try:
                    os.remove(html_path)
                except Exception:
                    pass
    if _HAS_REPORTLAB:
        try:
            pdf_fd, pdf_path = tempfile.mkstemp(suffix=".pdf")
            os.close(pdf_fd)
            c = canvas.Canvas(pdf_path, pagesize=letter)
            width, height = letter
            y = height - 72
            plain = re.sub(r'<[^>]+>', '', html_content)
            for chunk in plain.splitlines():
                while len(chunk) > 100:
                    c.drawString(72, y, chunk[:100])
                    chunk = chunk[100:]
                    y -= 14
                    if y < 72:
                        c.showPage()
                        y = height - 72
                c.drawString(72, y, chunk[:100])
                y -= 14
                if y < 72:
                    c.showPage()
                    y = height - 72
            c.save()
            return pdf_path
        except Exception as e:
            log_event("reportlab_failed", error=str(e))
            try:
                if pdf_path and os.path.exists(pdf_path):
                    os.remove(pdf_path)
            except Exception:
                pass
    log_event("pdf_generation_failed", user=username)
    return None

def export_docx(answer_text: str, risk_info: List[Dict[str,Any]], username: str = "anonymous") -> Optional[str]:
    if not _HAS_DOCX or DocxDocument is None:
        return None
    try:
        doc = DocxDocument()
        doc.add_heading('LawMate AI — Analysis', level=1)
        doc.add_paragraph(f'User: {username}    Date: {now_utc().isoformat()}')
        doc.add_heading('Answer', level=2)
        doc.add_paragraph(answer_text)
        doc.add_heading('Risk Summary', level=2)
        for r in risk_info:
            p = doc.add_paragraph()
            p.add_run(f"{r['level']}: ").bold = True
            p.add_run(r['clause'])
            if r.get('reasons'):
                p.add_run(" — " + ", ".join(r['reasons']))
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name
    except Exception as e:
        log_event("docx_failed", error=str(e))
        return None

# ---------- Session State ----------
if "law_db" not in st.session_state:
    st.session_state.law_db = LawDatabase()
if 'user' not in st.session_state:
    st.session_state.user = None
if 'session_token' not in st.session_state:
    st.session_state.session_token = None

# ---------- Sidebar Mode ----------
with st.sidebar:
    mode = st.radio("Mode", ["Normal", "Law Chatbot"])
st.session_state.mode = mode

# ---------- Law Chatbot UI ----------
if st.session_state.mode == "Law Chatbot":
    st.header("⚖️ Law Chatbot")
    query = st.text_input("Ask a legal question (Constitution / CrPC ):")
    
    if query:
        with st.spinner("Searching law database..."):
            try:
                safe_query = html.escape(query)
                results = st.session_state.law_db.query_laws(safe_query, top_k=3)

                if not results:
                    st.info("No results found.")
                else:
                    for r in results:
                        if r["type"] == "qa":
                            st.markdown(f"**Q:** {r['question']}")
                            st.markdown(f"**A:** {r['answer']}")
                            st.caption(f"Score: {r['score']:.2f}")
                        elif r["type"] == "ref":
                            st.markdown(f"**Reference:** [{r['title']}]({r['url']})")
                            st.caption(f"Score: {r['score']:.2f}")

            except Exception as e:
                st.error(f"Law Chatbot error: {e}")

# ---------- Normal UI ----------
if st.session_state.mode == "Normal":
    st.title("⚖️ LawMate AI")
    st.write("Welcome! Switch to 'Law Chatbot' from the sidebar to ask legal questions.")

# Sidebar: auth & settings + diagnostics + admin page toggle
with st.sidebar:
    st.header("Account & Settings")
    if st.session_state.user and is_session_expired():
        st.warning("Session expired due to inactivity.")
        st.session_state.user = None
        st.session_state.session_token = None
        st.rerun()

    if st.session_state.user is None:
        mode = st.radio("Login / Sign up / Reset", ["Login", "Sign up", "Reset password"], index=0)
        if mode == "Sign up":
            su_user = st.text_input("Username", key="su_user")
            su_pass = st.text_input("Password", type="password", key="su_pass")
            su_q = st.text_input("Security question (e.g. Your mother's maiden name)", key="su_q")
            su_a = st.text_input("Security answer", key="su_a")
            if st.button("Create account"):
                ok, msg = password_valid_policy(su_pass)
                if not ok:
                    st.error(msg)
                else:
                    with SessionLocal() as db:
                        if db.query(User).filter_by(username=su_user).first():
                            st.error("Username exists.")
                        else:
                            u = User(username=su_user, password_hash=hash_password(su_pass),
                                     sec_question=su_q.strip() if su_q else None,
                                     sec_answer_hash=hash_answer(su_a) if su_a else None)
                            db.add(u)
                            db.commit()
                            log_event("user_signup", user=su_user)
                            st.success("Account created. Please log in.")
        elif mode == "Reset password":
            rs_user = st.text_input("Username to reset", key="rs_user")
            if st.button("Fetch security question"):
                with SessionLocal() as db:
                    uu = db.query(User).filter_by(username=rs_user).first()
                    if not uu or not uu.sec_question:
                        st.error("User not found or no security question set.")
                    else:
                        st.session_state.reset_user = rs_user
                        st.session_state.reset_question = uu.sec_question
                        st.rerun()
            if st.session_state.get("reset_user") == rs_user and st.session_state.get("reset_question"):
                st.write("Security question:")
                st.write(st.session_state.reset_question)
                rs_answer = st.text_input("Answer", type="password",key="rs_ans")
                new_pw = st.text_input("New password", type="password", key="rs_newpw")
                if st.button("Reset password now"):
                    with SessionLocal() as db:
                        uu = db.query(User).filter_by(username=rs_user).first()
                        if not uu:
                            st.error("User disappeared.")
                        elif not uu.sec_answer_hash or not verify_answer(rs_answer, uu.sec_answer_hash):
                            st.error("Incorrect answer.")
                        else:
                            ok, msg = password_valid_policy(new_pw)
                            if not ok:
                                st.error(msg)
                            else:
                                uu.password_hash = hash_password(new_pw)
                                db.commit()
                                log_event("password_reset", user=rs_user)
                                st.success("Password reset. Please login.")
                                st.session_state.pop("reset_user", None)
                                st.session_state.pop("reset_question", None)
        else:
            li_user = st.text_input("Username", key="li_user")
            li_pwd = st.text_input("Password", type="password", key="li_pass")
            if st.button("Login"):
                with SessionLocal() as db:
                    dbu = db.query(User).filter_by(username=li_user).first()
                    if not dbu:
                        st.error("Invalid credentials.")
                        log_event("login_failed", user=li_user)
                    else:
                        if dbu.locked_until and dbu.locked_until > datetime.utcnow():
                            wait = int((dbu.locked_until - datetime.utcnow()).total_seconds() / 60) + 1
                            st.error(f"Account locked. Try again in {wait} minutes.")
                            log_event("login_locked", user=li_user)
                        elif verify_password(li_pwd, dbu.password_hash):
                            st.session_state.user = {"id": dbu.id, "username": dbu.username}
                            regenerate_session_token()
                            dbu.failed_logins = 0
                            dbu.locked_until = None
                            db.commit()
                            refresh_session_time()
                            log_event("login_success", user=li_user, session_id=st.session_state.session_token[:8])
                            st.success(f"Welcome, {dbu.username}!")
                            st.rerun()
                        else:
                            dbu.failed_logins = (dbu.failed_logins or 0) + 1
                            if dbu.failed_logins >= LOGIN_RATE_LIMIT:
                                dbu.locked_until = datetime.utcnow() + timedelta(minutes=LOGIN_LOCK_MIN)
                                st.error(f"Too many failed tries. Account locked for {LOGIN_LOCK_MIN} minutes.")
                                log_event("account_locked", user=li_user)
                            else:
                                st.error("Invalid credentials.")
                                log_event("login_failed", user=li_user)
                            db.commit()
    else:
        st.write(f"Logged in as **{st.session_state.user['username']}**")
        if st.button("Logout"):
            log_event("logout", user=st.session_state.user['username'], session=st.session_state.get("session_token"))
            st.session_state.user = None
            st.session_state.session_token = None
            st.rerun()

    st.markdown("---")
    st.header("App Settings")
    st.info(f"Upload : {MAX_UPLOAD_MB} MB | Session timeout: {SESSION_TIMEOUT_MIN} min")
    try:
        langs = tesseract_languages_available()
        st.write("Tesseract languages available:", langs if langs else "(none detected)")
    except Exception:
        st.write("Tesseract detection unavailable")

    st.markdown("---")
    st.subheader("Diagnostics / Admin")
    st.write("Quick checks")
    diag_cols = st.columns(3)
    with diag_cols[0]:
        st.write("OCR")
        st.write("✅" if _HAS_PYTESSERACT else "❌")
    with diag_cols[1]:
        st.write("HF Model")
        st.write("✅" if qa_pipeline else "❌")
    with diag_cols[2]:
        st.write("DOCX")
        st.write("✅" if _HAS_DOCX else "❌")
    if st.button("Open full Diagnostics page"):
        st.session_state.show_diagnostics_page = True

# If not logged in, stop main
if st.session_state.user is None:
    st.info("Log in or sign up to begin analyzing contracts.")
    st.stop()

# ---------- Diagnostics page (standalone) ----------
if st.session_state.get("show_diagnostics_page"):
    st.markdown("# Diagnostics")
    st.write("This page runs a series of checks and displays results below.")
    st.subheader("Environment & dependencies")
    env = {
        "python_platform": platform.platform(),
        "weasyprint": _HAS_WEASY,
        "reportlab": _HAS_REPORTLAB,
        "pytesseract": _HAS_PYTESSERACT,
        "pdf2image": _HAS_PDF2IMAGE,
        "transformers": _HAS_TRANSFORMERS,
        "docx": _HAS_DOCX,
        "argon2": _HAS_ARGON2,
    }
    st.json(env)

    st.subheader("DB connectivity test")
    try:
        with SessionLocal() as db:
            # count tables
            users_count = db.query(User).count()
            contracts_count = db.query(Contract).count()
            ocr_cache_count = db.query(OCRCache).count()
        st.success("DB connected")
        st.write({"users": users_count, "contracts": contracts_count, "ocr_cache": ocr_cache_count})
    except Exception as e:
        st.error("DB test failed: " + str(e))
        log_event("diag_db_failed", error=str(e))

    st.subheader("Model quick check")
    if qa_pipeline:
        try:
            st.write(str(qa_pipeline))
            st.success("QA pipeline available")
        except Exception as e:
            st.error("QA pipeline issue: " + str(e))
    else:
        st.warning("No QA pipeline loaded (transformers not available or model failed to load).")

    st.subheader("OCR quick test")
    try:
        sample_text = "Sample 123 ABC"
        if _HAS_PYTESSERACT:
            st.write("pytesseract available")
        else:
            st.warning("pytesseract not available")
    except Exception as e:
        st.error("OCR test failed: " + str(e))

    if st.button("Close diagnostics"):
        st.session_state.show_diagnostics_page = False
        st.rerun()

    st.stop()

# ---------- Main Area ----------
st.subheader("Upload / choose contract")
col1, col2 = st.columns([3,1])
with col1:
    uploaded = st.file_uploader("Upload PDF or TXT", type=["pdf","txt"], accept_multiple_files=False)
    # Batch OCR uploader
    batch_upload = st.file_uploader("Batch OCR: upload multiple PDFs/TXTs", type=["pdf","txt"], accept_multiple_files=True)
    sample = st.selectbox("Or sample", ["None", "Sample TXT", "Sample PDF"])
with col2:
    # Auto language detection option
    autodetect = st.checkbox("Auto-detect language for OCR", value=True)
    ocr_lang_choice = st.selectbox("OCR Language (fallback)", list(LANG_CODE_MAP.keys()))
    out_lang = st.selectbox("Translate output to", ["en", "hi", "fr", "es"], index=0)

# Batch OCR processing
if batch_upload:
    st.info(f"Processing {len(batch_upload)} files in parallel (Batch OCR).")
    # show progress and run parallel OCR
    results = []
    max_workers = min(4, len(batch_upload))
    with st.spinner("Running batch OCR..."):
        def process_file(up) -> Tuple[str, str, str]:
            name = sanitize_filename(up.name)
            raw = up.read()
            # decide language
            lang_code = LANG_CODE_MAP.get(ocr_lang_choice, "eng")
            if autodetect and _HAS_LANGDETECT:
                try:
                    guessed = langdetect_detect(raw[:4000].decode('utf-8', errors='ignore'))
                    if guessed.startswith("hi"):
                        lang_code = "hin"
                    elif guessed.startswith("fr"):
                        lang_code = "fra"
                    else:
                        lang_code = "eng"
                except Exception:
                    lang_code = LANG_CODE_MAP.get(ocr_lang_choice, "eng")
            text = ocr_cached(raw, lang_code)
            return (name, file_hash_bytes(raw), text)

        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = {ex.submit(process_file, up): up for up in batch_upload}
            for f in as_completed(futures):
                try:
                    name, fhash, text = f.result()
                    results.append({"file_name": name, "file_hash": fhash, "text_preview": text[:200], "ts": now_utc().isoformat()})
                except Exception as e:
                    results.append({"file_name": getattr(futures[f], "name", "unknown"), "file_hash": None, "text_preview": "", "error": str(e)})
    # display results
    st.subheader("Batch OCR results")
    df_batch = pd.DataFrame(results)
    st.dataframe(df_batch, use_container_width=True)
    # allow download per-item JSON
    for r in results:
        if r.get("file_hash"):
            key = r["file_hash"]
            payload = {"meta": {"user": st.session_state.user['username'], "ts": now_utc().isoformat()},
           "file_name": r["file_name"], "file_hash": r["file_hash"], "preview": r["text_preview"]}
            payload_bytes = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
            st.download_button(
                label=f"Download OCR JSON: {r['file_name']}",
                data=payload_bytes,
                file_name=f"ocr_{r['file_hash'][:8]}_{int(time.time())}.json",
                mime="application/json"
                )

    log_event("batch_ocr_completed", user=st.session_state.user['username'], count=len(results))

# Load prior file_data if present
text = ""
file_name = None
if 'file_data' in st.session_state:
    text = st.session_state.file_data.get('text', "")
    file_name = st.session_state.file_data.get('file_name')

# handle single file upload
if uploaded:
    if uploaded.size > MAX_UPLOAD_MB * 1024 * 1024:
        st.error(f"File too large (> {MAX_UPLOAD_MB} MB).")
        st.stop()
    raw = uploaded.read()
    file_name = sanitize_filename(uploaded.name)
    if uploaded.name.lower().endswith(".txt"):
        text = extract_txt(raw)
    else:
        lang_code = LANG_CODE_MAP.get(ocr_lang_choice, "eng")
        if autodetect and _HAS_LANGDETECT:
            try:
                guessed = langdetect_detect(raw[:4000].decode('utf-8', errors='ignore'))
                if guessed.startswith("hi"):
                    lang_code = "hin"
                elif guessed.startswith("fr"):
                    lang_code = "fra"
                else:
                    lang_code = "eng"
            except Exception:
                lang_code = LANG_CODE_MAP.get(ocr_lang_choice, "eng")
        text = ocr_cached(raw, lang_code)
    st.session_state.file_data = {"text": text, "file_name": file_name, "bytes_hash": file_hash_bytes(raw)}
    st.session_state.pop("last_result", None)
    refresh_session_time()
    log_event("file_uploaded", user=st.session_state.user['username'], file=file_name)
    

# sample selection
if sample != "None" and not uploaded and 'file_data' not in st.session_state:
    sample_path = "samples/sample.txt" if "TXT" in sample else "samples/sample.pdf"
    if os.path.exists(sample_path):
        with open(sample_path, "rb") as f:
            raw = f.read()
        if sample_path.endswith(".txt"):
            text = extract_txt(raw)
        else:
            text = ocr_cached(raw, LANG_CODE_MAP.get(ocr_lang_choice, "eng"))
        file_name = os.path.basename(sample_path)
        st.session_state.file_data = {"text": text, "file_name": file_name, "bytes_hash": file_hash_bytes(raw)}
        st.session_state.pop("last_result", None)
        
    else:
        st.warning("No sample files present; upload a file to test.")

if not text:
    st.info("Upload a file (PDF/TXT) or choose a sample.")
    st.stop()

# Document preview with highlights
st.markdown("---")
st.markdown("### Document preview (first 10k chars) — risky terms highlighted")
MAX_PREVIEW_CHARS = 10000
preview_text = text[:MAX_PREVIEW_CHARS] 
search_term = st.text_input("Search inside document (preview highlights):")
try:
    preview_html = render_preview_with_highlights(preview_text, search_term or "")
    st.markdown(preview_html, unsafe_allow_html=True)
except Exception as e:
    st.error("Preview rendering failed: " + str(e))


# Quick search in document

if search_term:
    try:
        highlighted_html = re.sub(
            re.escape(search_term),
            lambda m: f"<mark>{m.group(0)}</mark>",
            preview_html,
            flags=re.IGNORECASE
        )
        st.markdown("### Search results (preview)")
        st.markdown(highlighted_html, unsafe_allow_html=True)
    except Exception as e:
        st.error("Search highlighting failed: " + str(e))


# Q/A inputs
question = st.text_input("Ask a short, specific question about this contract:")
colA, colB = st.columns([1,1])
with colA:
    analyze = st.button("Analyze")
with colB:
    summary_btn = st.button("Quick Summary")

# Quick search in previous QAs (sidebar)
search_q = st.sidebar.text_input("Search previous QAs (keyword)", key="search_q")
st.sidebar.markdown("---")
# per-user stats
with SessionLocal() as db:
    uid = st.session_state.user['id']
    total_contracts = db.query(Contract).filter_by(owner_id=uid).count()
    total_qas = db.query(QACache).join(Contract, QACache.contract_id == Contract.id).filter(Contract.owner_id == uid).count()
    latest_qas = db.query(QACache).join(Contract, QACache.contract_id == Contract.id).filter(Contract.owner_id == uid).order_by(QACache.created_at.desc()).limit(20).all()
    avg_risk_score = None
    if latest_qas:
        tot = 0
        cnt = 0
        for q in latest_qas:
            try:
                info = json.loads(q.risk_info)
                for item in info:
                    lvl = item.get("level")
                    val = {"Low": 0, "Medium": 1, "High": 2}.get(lvl, 0)
                    tot += val
                    cnt += 1
            except Exception:
                continue
        avg_risk_score = (tot / cnt) if cnt else 0

st.sidebar.markdown("### Your stats")
st.sidebar.write(f"Contracts: **{total_contracts}**")
st.sidebar.write(f"Q&As: **{total_qas}**")
if avg_risk_score is not None:
    st.sidebar.write(f"Avg risk (0 low — 2 high): **{avg_risk_score:.2f}**")
st.sidebar.markdown("---")

# handle analyze/summary
if analyze:
    if not question:
        st.error("Please enter a question first.")
    else:
        st.session_state.processing = "analyze"
        st.rerun()

if summary_btn:
    st.session_state.processing = "summary"
    st.rerun()

# Processing analyze
if st.session_state.get("processing") == "analyze":
    with st.spinner("Analyzing with model / heuristics..."):
        try:
            refresh_session_time()
            # store contract (versioning)
            with SessionLocal() as db:
                owner_id = st.session_state.user['id']
                prev = db.query(Contract).filter_by(file_name=file_name, owner_id=owner_id).order_by(Contract.created_at.desc()).first()
                if prev:
                    c = Contract(file_name=file_name, text=text, owner_id=owner_id, parent_id=prev.id)
                else:
                    c = Contract(file_name=file_name, text=text, owner_id=owner_id)
                db.add(c)
                db.commit()
                contract_id = c.id
            # query model
            answer = query_model_with_chunking(text, question)
            answer_translated = translate_text(answer, dest=out_lang)
            # 🔹 Find relevant laws from your datasets
            relevant_laws = st.session_state.law_db.query_laws(
                question + " " + text[:1000],
                top_k=3
                )

            # extract clauses
            clauses = [s.strip() for s in re.split(r'\n{1,}|(?<=[.?!])\s+', answer_translated) if s.strip()][:100]
            risk_info = []
            for cl in clauses:
                r = compute_risk(cl)
                law_refs = st.session_state.law_db.query_laws(cl, top_k=1)
                law_ref = law_refs[0] if law_refs else None
                risk_info.append({"clause": cl, "level": r["level"], "score": r["score"], "reasons": r["reasons"],"law_ref": law_ref})
            # save QA entry
            with SessionLocal() as db:
                qa = QACache(contract_id=contract_id, question=question, answer=answer_translated, risk_info=json.dumps(risk_info))
                db.add(qa)
                db.commit()
                qa_id = qa.id
            st.session_state.last_result = {
                 "answer": answer_translated,
                 "risk_info": risk_info,
                 "contract_id": contract_id,
                 "question": question,
                 "qa_id": qa_id,
                 "relevant_laws": relevant_laws  
                 }

            log_event("analysis_run", user=st.session_state.user['username'], contract_id=contract_id, qa_id=qa_id)
            st.success("Analysis complete.")
        except Exception as e:
            st.error(f"Analysis failed: {e}")
            log_event("analysis_failed", user=st.session_state.user['username'], error=str(e))
    st.session_state.processing = None
    st.rerun()

# Processing summary
elif st.session_state.get("processing") == "summary":
    with st.spinner("Generating summary..."):
        try:
            refresh_session_time()
            if qa_pipeline:
                try:
                    out = qa_pipeline(question="Summarize the document", context=text[:3000]) if getattr(qa_pipeline, "task", "") == "question-answering" else qa_pipeline("Summarize:\n" + text[:3000], max_length=200)
                    summary_text = ""
                    if isinstance(out, dict) and "answer" in out:
                        summary_text = out.get("answer", "")
                    elif isinstance(out, list) and out and isinstance(out[0], dict):
                        summary_text = out[0].get("generated_text", "") or out[0].get("answer", "")
                    else:
                        summary_text = str(out)
                except Exception:
                    summary_text = text[:500]
            else:
                summary_text = text[:500]
            st.session_state.summary_temp = summary_text
            log_event("summary_generated", user=st.session_state.user['username'])
        except Exception as e:
            st.error(f"Summary failed: {e}")
            log_event("summary_failed", user=st.session_state.user['username'], error=str(e))
    st.session_state.processing = None
    st.rerun()

# show summary
if st.session_state.get("summary_temp"):
    st.markdown("---")
    st.subheader("Auto Summary")
    st.write(st.session_state.pop("summary_temp"))

# show last result
if st.session_state.get("last_result"):
    res = st.session_state.last_result
    st.markdown("---")
    st.subheader("Analysis Result")
    st.markdown("**Answer:**")
    st.write(res["answer"])
    st.markdown("**Clause-level Risk Analysis**")
    # Show relevant laws
    if st.session_state.last_result.get("relevant_laws"):
        st.subheader("Relevant Laws / References")
        for law in st.session_state.last_result["relevant_laws"]:
            if law["type"] == "qa":
                st.markdown(f"**QA:** {law['question']} → {law['answer']}")
            elif law["type"] == "ref":
                st.markdown(f"**Law:** [{law['title']}]({law['url']})")


    df_rows = []
    for i, r in enumerate(res["risk_info"]):
        df_rows.append({"Clause #": i+1, "Risk Level": r["level"], "Score": r.get("score", 0), "Clause": r["clause"]})
    df_risk = pd.DataFrame(df_rows)

    if not df_risk.empty:
        try:
            styled = df_risk.style.background_gradient(subset=["Score"], cmap="Reds")
            st.dataframe(styled, use_container_width=True)
        except Exception:
            st.dataframe(df_risk, use_container_width=True)

    for i, r in enumerate(res["risk_info"]):
        color = {"High": "#ff4b4b", "Medium": "#ffb04b", "Low": "#7bd389"}.get(r["level"], "#cccccc")
        with st.expander(f"Clause {i+1} — Risk: {r['level']} (score {r['score']})"):
            st.markdown(f"<div style='color:{color}'><strong>Clause:</strong> {r['clause']}</div>", unsafe_allow_html=True)
            if r.get('reasons'):
                st.write("Reasons:")
                for reason in r['reasons']:
                    st.write("- " + reason)
                    if r.get("law_ref"):
                        law = r["law_ref"]
                        st.write("⚖️ Related Law:")
                        if law["type"] == "ref":
                            st.markdown(f"- [{law['title']}]({law['url']})")
                        elif law["type"] == "qa":
                            st.markdown(f"- QA: {law['question']} → {law['answer']}")


    if not df_risk.empty:
        risk_map = {"Low": 0, "Medium": 1, "High": 2}
        heat_vals = [risk_map.get(l, 0) for l in df_risk['Risk Level']]
        fig = px.imshow([heat_vals], labels=dict(x="Clause position", y="", color="Risk score"), x=list(range(1, len(heat_vals)+1)), y=["Document"])
        st.plotly_chart(fig, use_container_width=True)

    st.markdown("### Risk Distribution")
    if not df_risk.empty:
        counts = df_risk['Risk Level'].value_counts().reindex(['High', 'Medium', 'Low']).fillna(0)
        fig2 = px.pie(values=counts.values, names=counts.index, title="Risk Level Distribution")
        st.plotly_chart(fig2, use_container_width=True)

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        if st.button("Download PDF Report"):
            try:
                pdfp = export_pdf_with_metadata(res["answer"], res["risk_info"], username=st.session_state.user['username'])
                if pdfp and os.path.exists(pdfp):
                    with open(pdfp, "rb") as f:
                        st.download_button("Download PDF", f, file_name=f"analysis_{int(time.time())}.pdf", mime="application/pdf")
                    try:
                        os.remove(pdfp)
                    except Exception:
                        pass
                    log_event("export_pdf", user=st.session_state.user['username'])
                else:
                    st.error("PDF generation not available on this system.")
            except Exception as e:
                st.error(f"PDF export failed: {e}")
    with c2:
        if _HAS_DOCX and st.button("Download DOCX Report"):
            try:
                docx_p = export_docx(res["answer"], res["risk_info"], username=st.session_state.user['username'])
                if docx_p and os.path.exists(docx_p):
                    with open(docx_p, "rb") as f:
                        st.download_button("Download DOCX", f, file_name=f"analysis_{int(time.time())}.docx")
                    try:
                        os.remove(docx_p)
                    except Exception:
                        pass
                    log_event("export_docx", user=st.session_state.user['username'])
                else:
                    st.error("DOCX generation not available.")
            except Exception as e:
                st.error(f"DOCX export failed: {e}")
    with c3:
        if st.button("Download JSON"):
            payload = {
                "meta": {"user": st.session_state.user['username'], "ts": now_utc().isoformat()},
                "file_name": res.get("file_name", "unknown"),
                "file_hash": res.get("file_hash", "na"),
                "preview": res.get("answer", "")[:200]
                }
            payload_bytes = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
            st.download_button(
                label=f"Download OCR JSON: {payload['file_name']}",
                data=payload_bytes,
                file_name=f"ocr_{payload['file_hash'][:8]}_{int(time.time())}.json",
                mime="application/json"
                )
            log_event("export_json", user=st.session_state.user['username'])

    with c4:
        if st.button("🔊 Read Aloud"):
            try:
                if platform.system() == 'Darwin':
                    subprocess.run(["say", res["answer"]])
                else:
                    import pyttsx3
                    engine = pyttsx3.init()
                    engine.say(res["answer"])
                    engine.runAndWait()
                log_event("tts_played", user=st.session_state.user['username'])
            except Exception as e:
                st.error("TTS failed: " + str(e))

    st.markdown("### Actions")
    if st.button("Diff with new upload"):
        new_upload = st.file_uploader("Upload new version to diff (pdf/txt)", type=["pdf","txt"], key="diff_upload")
        if new_upload:
            new_raw = new_upload.read()
            if new_upload.name.lower().endswith(".txt"):
                new_text = extract_txt(new_raw)
            else:
                new_text = ocr_cached(new_raw, LANG_CODE_MAP.get(ocr_lang_choice, "eng"))
            diff = difflib.unified_diff(res["answer"].splitlines(), new_text.splitlines(), lineterm="")
            diff_text = "\n".join(diff)
            st.text_area("Diff (unified)", value=diff_text[:20000], height=400)
            log_event("diff_generated", user=st.session_state.user['username'])

    if st.button("Delete this QA"):
        st.session_state.confirm_delete_qa = True

    if st.session_state.get('confirm_delete_qa'):
        st.warning("Confirm delete of this QA?")
        c_yes, c_no = st.columns(2)
        with c_yes:
            if st.button("Yes, delete now"):
                try:
                    qa_id = res.get("qa_id")
                    if qa_id:
                        with SessionLocal() as db:
                            db.query(QACache).filter_by(id=qa_id).delete()
                            db.commit()
                    st.session_state.pop("last_result", None)
                    st.session_state.pop("confirm_delete_qa", None)
                    log_event("qa_deleted", user=st.session_state.user['username'], qa_id=qa_id)
                    st.success("Deleted.")
                except Exception as e:
                    st.error(f"Delete failed: {e}")
                    log_event("qa_delete_failed", user=st.session_state.user['username'], error=str(e))
                st.rerun()
        with c_no:
            if st.button("Cancel"):
                st.session_state.pop("confirm_delete_qa", None)
                st.rerun()

# Sidebar: previous contracts & QAs (lazy load + search)
st.sidebar.markdown("---")
st.sidebar.header("Your Contracts & Q&A (lazy)")
with SessionLocal() as db:
    uid = st.session_state.user['id']
    contracts = db.query(Contract).filter_by(owner_id=uid).order_by(Contract.created_at.desc()).all()
    for c in contracts:
        if st.sidebar.checkbox(f"Show: {c.file_name} — {c.created_at.date()}", key=f"show_contract_{c.id}"):
            st.sidebar.write(c.text[:500] + ("..." if len(c.text) > 500 else ""))
            qas = db.query(QACache).filter_by(contract_id=c.id).order_by(QACache.created_at.desc()).all()
            for q in qas:
                if search_q:
                    if search_q.lower() not in (q.question or "").lower() and search_q.lower() not in (q.answer or "").lower():
                        continue
                if st.sidebar.button(f"Load QA {q.id}", key=f"qa_{q.id}"):
                    try:
                        risk_info = json.loads(q.risk_info) if q.risk_info else []
                    except Exception:
                        risk_info = []
                    st.session_state.last_result = {"answer": q.answer, "risk_info": risk_info, "contract_id": c.id, "question": q.question, "qa_id": q.id}
                    st.rerun()

            if st.sidebar.button(f"Delete contract {c.id}", key=f"delc_{c.id}"):
                st.session_state[f"confirm_delete_contract_{c.id}"] = True
            if st.session_state.get(f"confirm_delete_contract_{c.id}"):
                st.sidebar.warning("Confirm delete (all QAs will be removed)")
                if st.sidebar.button(f"Confirm delete {c.id}", key=f"confirm_del_{c.id}"):
                    try:
                        with SessionLocal() as db_del:
                            db_del.execute(sa_delete(Contract).where(Contract.id == c.id))
                            db_del.commit()
                        st.sidebar.success("Deleted")
                        log_event("contract_deleted", user=st.session_state.user['username'], contract_id=c.id)
                    except Exception as e:
                        st.sidebar.error("Delete failed")
                        log_event("contract_delete_failed", user=st.session_state.user['username'], error=str(e))
                    st.rerun()
